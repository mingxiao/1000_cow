#!/usr/bin/python3
from docx import Document
import csv
import os

def is_bold(cell):
  if cell.paragraphs[0].runs[0].bold:
    return 'true'
  else:
   return 'false'

def get_row_output(row, include_top_30_check=False):
  result =  [
    row.cells[0].text.strip(),
    row.cells[1].text.strip(),
  ]
  if include_top_30_check:
    cell_is_bold = is_bold(row.cells[0])
    result.append(cell_is_bold)
  return result

def process_file(filename, include_top_30_check=False):
  # processes files in the raw/ directory and outputs a csv file in the processed/ directory
  file_path = os.path.join(os.getcwd(), 'raw', filename)
  assert os.path.exists(file_path)
  basename = os.path.splitext(filename)[0]
  output_file_path = os.path.join(os.getcwd(), 'processed', f'{basename}.csv')
  doc = Document(file_path)
  assert doc
  data_table = doc.tables[0]
  assert data_table
  header = data_table.rows[0]
  table_content = data_table.rows[1:]
  csv_header = [header.cells[0].text, header.cells[1].text]
  if include_top_30_check:
    csv_header.append('In top 30')
  with open(output_file_path, 'w') as csvfile:
    filewriter = csv.writer(csvfile, delimiter=',',
                            quotechar='|', quoting=csv.QUOTE_MINIMAL)
    filewriter.writerow(csv_header)
    for row in table_content:
      row_output = get_row_output(row, include_top_30_check)
      filewriter.writerow(row_output)



# file_path = os.path.join(os.getcwd(), 'raw', 'aav8391_Data_S2.docx')
# document = Document(file_path)
# >>> doc.tables[0].rows[0].cells[0].text
# 'Taxonomic group'
# >>> doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0].bold
# print(document)
# process_file('aav8391_Data_S2.docx', include_top_30_check=True)
# process_file('aav8391_Data_S3.docx', include_top_30_check=False)
process_file('aav8391_Data_S4.docx', include_top_30_check=False)