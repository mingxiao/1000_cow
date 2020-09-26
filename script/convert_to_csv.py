#!/usr/bin/python3
from docx import Document
import csv
import os

DATA_DIR=os.path.join(os.path.dirname(__file__), '..', '1000_cow_data')

def is_bold(cell):
  if len(cell.paragraphs) == 0:
    return False
  if len(cell.paragraphs[0].runs) == 0:
    return False
  if cell.paragraphs[0].runs[0].bold:
    return True
  else:
   return False

def get_row_output(row, include_top_30_check=False):
  assert len(row.cells) >= 2
  result =  [
    row.cells[0].text.strip(),
    row.cells[1].text.strip(),
  ]
  if include_top_30_check:
    cell_is_bold = is_bold(row.cells[0])
    if cell_is_bold:
      result.append('true')
    else:
      result.append('false')
  return result

def convert_docx_table_to_csv(filename, include_top_30_check=False):
  # processes files in the raw/ directory and outputs a csv file in the processed/ directory
  file_path = os.path.join(DATA_DIR, 'raw', filename)
  assert os.path.exists(file_path)
  basename = os.path.splitext(filename)[0]
  output_file_path = os.path.join(DATA_DIR, 'processed', f'{basename}.csv')
  doc = Document(file_path)
  assert doc
  data_table = doc.tables[0]
  assert data_table
  header = data_table.rows[0]
  table_content = data_table.rows[1:]
  csv_header = [header.cells[0].text, header.cells[1].text]
  if include_top_30_check:
    csv_header.append('In top 30')
  with open(output_file_path, 'w') as csvfile:
    filewriter = csv.writer(csvfile, delimiter=',',
                            quotechar='|', quoting=csv.QUOTE_MINIMAL)
    filewriter.writerow(csv_header)
    for row in table_content:
      row_output = get_row_output(row, include_top_30_check)
      filewriter.writerow(row_output)

if __name__ == '__main__':
  print("Converting docx to csv files")
  convert_docx_table_to_csv('aav8391_Data_S2.docx', include_top_30_check=True)
  convert_docx_table_to_csv('aav8391_Data_S3.docx', include_top_30_check=False)
  convert_docx_table_to_csv('aav8391_Data_S4.docx', include_top_30_check=False)