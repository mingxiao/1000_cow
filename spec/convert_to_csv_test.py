import pytest
import docx
import sys
import os
src_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'script')
sys.path.append(src_dir)
import convert_to_csv

class TestConvertToCSV:
  @pytest.fixture
  def table(self):
    document = docx.Document()
    table = document.add_table(rows=1, cols=3)
    return table

  @pytest.fixture
  def row(self, table):
    row = table.rows[0]
    return row

  def test_is_bold_True(self, table):
    test_cell = table.rows[0].cells[0]
    test_cell.add_paragraph('')
    test_cell.paragraphs[0].add_run('bold').bold = True
    assert convert_to_csv.is_bold(test_cell) == True

  def test_is_bold_False(self, table):
    test_cell = table.rows[0].cells[0]
    assert convert_to_csv.is_bold(test_cell) == False
    test_cell.add_paragraph('')
    test_cell.paragraphs[0].add_run('word')
    assert convert_to_csv.is_bold(test_cell) == False

  def test_get_row_output(self, row):
    row.cells[0].add_paragraph('cell1  ')
    row.cells[1].add_paragraph('cell2\t')
    result = convert_to_csv.get_row_output(row)
    assert result == ['cell1', 'cell2']

    result = convert_to_csv.get_row_output(row, include_top_30_check=True)
    assert result == ['cell1', 'cell2',  'false']

  def test_convert_docx_table_to_csv(self):
    return False